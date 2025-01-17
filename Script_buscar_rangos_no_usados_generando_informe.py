import pandas as pd
from tkinter import Tk
from tkinter.filedialog import askopenfilename
from docx import Document

#Creo una ventana oculta de tkinter
Tk().withdraw()

#Abro la ventana del explorador de archivos donde puedo seleccionar el documento CSV
ruta_archivo = askopenfilename(
    title="Selecciona el archivo CSV",
    filetypes=[("Archivos CSV", "*.csv"), ("Todos los archivos", "*.*")]
)

#Creo el documento donde se guardara el informe
doc = Document()
doc.add_heading("Reporte de Analisis de valores repetido y rangos disponibles", 0)

#Variables
#Cambiamos el valor de columna de acuerdo a la columna que esperamos analizar
columnaAnalisis = "LumNumSer"

#Validacion de que el usuario si selecciono el archivo
if ruta_archivo:
    print(f"Archivo seleccionado: {ruta_archivo}")
    doc.add_paragraph(f"Archivo seleccionado: {ruta_archivo}")
    
    #Lectura del archivo CSV con pandas
    csv = pd.read_csv(ruta_archivo, delimiter=";")
    
    if columnaAnalisis in csv.columns:
        columna = csv[columnaAnalisis]
        
        if not columna.empty:
            print(f"La columna que vas a manejar es: {columna.name} y tiene los siguientes valores iniciales \n{columna.head()}")
            doc.add_paragraph(f"La columna que vas a manejar es: {columna.name} y tiene los siguientes valores iniciales:")
            doc.add_paragraph(f"{columna.head()}")
            
        else:
            print("La columna no posee valores.")
            doc.add_paragraph("La columna no posee valores.")
        
        #Si todo sale bien hasta este punto procederemos a verificar en que parte de la columna existen valores repetidos
        valoresDuplicados = columna[columna.duplicated()].unique()
        
        if len(valoresDuplicados) > 0:
            print(f"Valores duplicados encontrados: {valoresDuplicados}")
            doc.add_paragraph(f"Valores duplicados encontrados: {valoresDuplicados}")
        else:
            print("No se encontraron valores duplicados.")
            doc.add_paragraph("No se encontraron valores duplicados.")
        
        #Si hay duplicados es de suma importancia saber la concurrencia del error
        conteoValores = columna.value_counts()
        conteoDuplicados = conteoValores[conteoValores > 1]
        
        if not conteoDuplicados.empty:
            print("\nConteo de cada valor en la columna:")
            print(conteoDuplicados)
            doc.add_paragraph("\nConteo de cada valor en la columna:")
            doc.add_paragraph(f"{conteoDuplicados}")
            
            #Frecuencia de las repeticiones
            frecuenciaDuplicado = conteoDuplicados.value_counts()
            print("\nFrecuencia de repeticiones:")
            print(frecuenciaDuplicado)
            doc.add_paragraph("\nFrecuencia de repeticiones:")
            doc.add_paragraph(f"{frecuenciaDuplicado}")
        else:
            print("No se encontraron duplicados para contar")
            doc.add_paragraph("No se encontraron duplicados para contar.")
    else:
        print("No existe la columna que esperas analizar")
        doc.add_paragraph("No existe la columna que esperas analizar.")
else:
    print("No se ha seleccionado ningun archivo.")
    doc.add_paragraph("No se ha seleccionado ningún archivo.")

#Guardamos el Reporte
reporte_path = 'reporte_valores_repetidos.docx'
doc.save(reporte_path)

print(f"Reporte generado y guardado en: {reporte_path}")
    

